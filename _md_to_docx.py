import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── colours ──────────────────────────────────────────────────────────────────
C_ACCENT  = RGBColor(0x2E, 0x86, 0xAB)   # steel blue
C_DARK    = RGBColor(0x1A, 0x1A, 0x2E)   # near-black
C_MUTED   = RGBColor(0x55, 0x55, 0x66)   # muted grey
C_BG_TBL  = RGBColor(0xEE, 0xF4, 0xF9)   # light blue table header bg
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_QUOTE   = RGBColor(0x44, 0x77, 0x9A)   # blockquote text


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **borders):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, color in borders.items():
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_run_bold(para, text: str, size: int = 11, color: RGBColor = None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


# ── markdown helpers ──────────────────────────────────────────────────────────

def parse_inline(para, text: str, base_size: int = 11, base_color: RGBColor = None):
    """Add runs to *para* handling **bold** and *italic* inline markdown."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()])
            r.font.size = Pt(base_size)
            if base_color:
                r.font.color.rgb = base_color
        if m.group(1).startswith('**'):
            r = para.add_run(m.group(2))
            r.bold = True
            r.font.size = Pt(base_size)
            if base_color:
                r.font.color.rgb = base_color
        elif m.group(1).startswith('*'):
            r = para.add_run(m.group(3))
            r.italic = True
            r.font.size = Pt(base_size)
            if base_color:
                r.font.color.rgb = base_color
        else:  # backtick code
            r = para.add_run(m.group(4))
            r.font.name = 'Courier New'
            r.font.size = Pt(base_size - 1)
            if base_color:
                r.font.color.rgb = base_color
        pos = m.end()
    if pos < len(text):
        r = para.add_run(text[pos:])
        r.font.size = Pt(base_size)
        if base_color:
            r.font.color.rgb = base_color


# ── main builder ──────────────────────────────────────────────────────────────

def convert(md_path: str, out_path: str):
    doc = Document()

    # page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip('\n')

        # ── fenced code block (mermaid / unicode art) ──────────────────────
        if raw.strip().startswith('```'):
            lang = raw.strip()[3:].strip()
            i += 1
            block_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                block_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # skip closing ```

            if lang.lower() == 'mermaid':
                # render as a shaded info box
                p = doc.add_paragraph()
                p.paragraph_format.left_indent  = Cm(0.8)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                label = p.add_run('📊  [Diagram — see rendered Markdown for flowchart/chart]')
                label.italic = True
                label.font.size = Pt(10)
                label.font.color.rgb = C_ACCENT
            else:
                # unicode art / text block — monospace box
                for bl in block_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent  = Cm(0.8)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after  = Pt(0)
                    r = p.add_run(bl if bl else ' ')
                    r.font.name = 'Courier New'
                    r.font.size = Pt(9)
                    r.font.color.rgb = C_DARK
                doc.add_paragraph()  # spacer
            continue

        # ── horizontal rule ────────────────────────────────────────────────
        if re.match(r'^-{3,}$', raw.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), '2E86AB')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── headings ──────────────────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', raw)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
            p.paragraph_format.space_after  = Pt(4)
            if level == 1:
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(22)
                r.font.color.rgb = C_ACCENT
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(16)
                r.font.color.rgb = C_DARK
            elif level == 3:
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(13)
                r.font.color.rgb = C_ACCENT
            else:
                r = p.add_run(text)
                r.bold = True
                r.italic = True
                r.font.size = Pt(11)
                r.font.color.rgb = C_MUTED
            i += 1
            continue

        # ── blockquote ────────────────────────────────────────────────────
        if raw.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1.0)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            parse_inline(p, raw[2:], base_size=11, base_color=C_QUOTE)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # ── table ─────────────────────────────────────────────────────────
        if raw.startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].startswith('|'):
                row_raw = lines[i].strip().strip('|')
                cells = [c.strip() for c in row_raw.split('|')]
                table_rows.append(cells)
                i += 1
            # remove separator row (---)
            table_rows = [r for r in table_rows if not all(re.match(r'^[-:]+$', c) for c in r if c)]
            if not table_rows:
                continue

            ncols = max(len(r) for r in table_rows)
            tbl = doc.add_table(rows=len(table_rows), cols=ncols)
            tbl.style = 'Table Grid'

            for ri, row in enumerate(table_rows):
                for ci, cell_text in enumerate(row):
                    if ci >= ncols:
                        break
                    cell = tbl.cell(ri, ci)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after  = Pt(3)
                    if ri == 0:
                        set_cell_bg(cell, '2E86AB')
                        r = p.add_run(re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text))
                        r.bold = True
                        r.font.size = Pt(10)
                        r.font.color.rgb = C_WHITE
                    else:
                        bg = 'EEF4F9' if ri % 2 == 0 else 'FFFFFF'
                        set_cell_bg(cell, bg)
                        parse_inline(p, cell_text, base_size=10)

            doc.add_paragraph()  # spacer after table
            continue

        # ── bullet list ───────────────────────────────────────────────────
        m = re.match(r'^(\s*)[-*]\s+(.*)', raw)
        if m:
            indent = len(m.group(1))
            text   = m.group(2)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent  = Cm(0.5 + indent * 0.3)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            parse_inline(p, text, base_size=11)
            i += 1
            continue

        # ── numbered list ─────────────────────────────────────────────────
        m = re.match(r'^\d+\.\s+(.*)', raw)
        if m:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            parse_inline(p, m.group(1), base_size=11)
            i += 1
            continue

        # ── blank line ────────────────────────────────────────────────────
        if raw.strip() == '':
            i += 1
            continue

        # ── plain paragraph ───────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(6)
        parse_inline(p, raw, base_size=11)
        i += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    base = r'c:\Users\tzvia\ds-course\projects\DS-ml-project'
    convert(f'{base}\\PROJECT_STORY.md', f'{base}\\PROJECT_STORY.docx')
